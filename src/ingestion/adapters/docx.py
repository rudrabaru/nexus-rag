import docx
import hashlib
import logging
from datetime import datetime
from src.crawling.metadata import CrawledDocument, AdapterResult, VisualChunkDraft
from src.ingestion.base import IngestionAdapter

logger = logging.getLogger(__name__)


class DocxAdapter(IngestionAdapter):
    """
    Adapter for reading local Word Documents (.docx) and extracting text, tables, and images.
    """

    async def ingest(
        self, source: str, extract_visuals: bool = False, **kwargs
    ) -> AdapterResult:
        logger.info(f"DocxAdapter reading: {source}, extract_visuals={extract_visuals}")

        local_path = source
        _temp_file = None

        if source.startswith("http://") or source.startswith("https://"):
            import tempfile
            import urllib.request
            import os

            try:
                _temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                headers = {"User-Agent": "Mozilla/5.0"}
                req = urllib.request.Request(source, headers=headers)
                with urllib.request.urlopen(req, timeout=30) as resp:
                    _temp_file.write(resp.read())
                _temp_file.flush()
                local_path = _temp_file.name
            except Exception as e:
                logger.error(f"Failed to download DOCX {source}: {e}")
                if _temp_file:
                    os.unlink(_temp_file.name)
                return AdapterResult(documents=[], visual_chunks=[])

        try:
            doc = docx.Document(local_path)
            text_blocks = []
            visual_chunks = []

            doc_id = hashlib.md5(source.encode()).hexdigest()

            processor = None
            asset_store = None
            if extract_visuals:
                try:
                    from src.ingestion.visual_processor import VisualProcessor
                    from src.ingestion.asset_store import LocalAssetStore

                    processor = VisualProcessor()
                    asset_store = LocalAssetStore()
                except Exception as e:
                    logger.error(f"Failed to load visual extraction dependencies: {e}")
                    extract_visuals = False

            # Iterate through block-level elements in document order
            # docx-python's doc.paragraphs and doc.tables don't preserve relative order easily out of the box,
            # but we can try our best. For a perfect block-level iteration, we'd need to parse XML directly.
            # A common approach is to iterate over doc.element.body
            for element in doc.element.body:
                if element.tag.endswith("p"):
                    # Paragraph
                    para = docx.text.paragraph.Paragraph(element, doc)
                    if para.text.strip():
                        style_name = para.style.name.lower() if para.style else "normal"
                        if "header" in style_name or "footer" in style_name:
                            continue

                        if style_name.startswith("heading"):
                            level = style_name.replace("heading", "").strip()
                            try:
                                level_int = int(level)
                                prefix = "#" * level_int
                                text_blocks.append(f"{prefix} {para.text}\n\n")
                            except ValueError:
                                text_blocks.append(f"# {para.text}\n\n")
                        else:
                            text_blocks.append(f"{para.text}\n\n")

                    # Extract inline images and charts (including SmartArt)
                    if extract_visuals and processor and asset_store:
                        # Find all blip elements (both w:drawing and mc:AlternateContent)
                        blips = para._p.xpath(".//a:blip")
                        for blip in blips:
                            rel_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if rel_id and rel_id in doc.part.rels:
                                try:
                                    rel = doc.part.rels[rel_id]
                                    img_data = rel.target_part.blob
                                    img_index = len(visual_chunks)
                                    asset_ref = asset_store.save(
                                        doc_id, img_data, f"img_{img_index}.png"
                                    )
                                    desc = processor.describe_image(img_data)

                                    visual_chunks.append(
                                        VisualChunkDraft(
                                            text=desc,
                                            asset_ref=asset_ref,
                                            asset_type="photo",
                                        )
                                    )
                                    text_blocks.append(
                                        f"\n\n> **Visual Element:**\n> {desc}\n\n"
                                    )
                                except Exception as e:
                                    logger.error(
                                        f"Failed to process inline image {rel_id}: {e}"
                                    )

                elif element.tag.endswith("tbl"):
                    # Table
                    table = docx.table.Table(element, doc)
                    markdown_table = []
                    for i, row in enumerate(table.rows):
                        row_data = [
                            cell.text.replace("\n", " ").strip() for cell in row.cells
                        ]
                        markdown_table.append("| " + " | ".join(row_data) + " |")
                        if i == 0:
                            # Add separator
                            markdown_table.append(
                                "|" + "|".join(["---" for _ in row_data]) + "|"
                            )
                    if markdown_table:
                        text_blocks.append("\n".join(markdown_table) + "\n\n")

            markdown_content = "".join(text_blocks)
            word_count = len(markdown_content.split())

            title = source.split("/")[-1].split("\\")[-1]

            crawled_doc = CrawledDocument(
                url=f"file:///{source}",
                title=title,
                markdown_content=markdown_content,
                crawl_depth=0,
                crawled_at=datetime.utcnow(),
                word_count=word_count,
                status_code=200,
            )

            return AdapterResult(documents=[crawled_doc], visual_chunks=visual_chunks)

        except Exception as e:
            logger.error(f"Failed to read DOCX {source}: {e}")
            return AdapterResult(documents=[], visual_chunks=[])
        finally:
            if _temp_file:
                import os

                try:
                    os.unlink(_temp_file.name)
                except Exception as cleanup_error:
                    logger.warning(
                        f"Failed to clean up temp file {_temp_file.name}: {cleanup_error}"
                    )
